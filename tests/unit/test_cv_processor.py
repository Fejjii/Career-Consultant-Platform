"""Tests for the CV processing pipeline."""

from __future__ import annotations

import io

import pytest
from docx import Document

from career_intel.services.cv_processor import (
    CVProcessingError,
    clean_cv_text,
    extract_text_from_bytes,
    process_cv,
    truncate_to_token_limit,
    validate_cv_upload,
)


class TestValidateCvUpload:
    def test_valid_txt(self) -> None:
        validate_cv_upload(b"some content", "cv.txt")

    def test_valid_pdf(self) -> None:
        validate_cv_upload(b"%PDF-1.7\nfake-pdf", "resume.pdf")

    def test_fake_pdf_rejected(self) -> None:
        with pytest.raises(CVProcessingError, match="does not match its extension"):
            validate_cv_upload(b"not-a-real-pdf", "resume.pdf")

    def test_fake_docx_rejected(self) -> None:
        with pytest.raises(CVProcessingError, match="DOCX appears malformed"):
            validate_cv_upload(b"PK-not-a-docx", "resume.docx")

    def test_binary_txt_rejected(self) -> None:
        with pytest.raises(CVProcessingError, match="binary content"):
            validate_cv_upload(b"hello\x00world", "resume.txt")

    def test_empty_filename_rejected(self) -> None:
        with pytest.raises(CVProcessingError, match="Filename is required"):
            validate_cv_upload(b"data", "")

    def test_unsupported_extension_rejected(self) -> None:
        with pytest.raises(CVProcessingError, match="Unsupported file type"):
            validate_cv_upload(b"data", "cv.xlsx")

    def test_empty_data_rejected(self) -> None:
        with pytest.raises(CVProcessingError, match="empty"):
            validate_cv_upload(b"", "cv.txt")

    def test_oversized_file_rejected(self) -> None:
        big_data = b"x" * (6 * 1024 * 1024)
        with pytest.raises(CVProcessingError, match="too large"):
            validate_cv_upload(big_data, "cv.txt", max_file_bytes=5 * 1024 * 1024)

    def test_custom_size_limit(self) -> None:
        data = b"x" * 1000
        with pytest.raises(CVProcessingError, match="too large"):
            validate_cv_upload(data, "cv.txt", max_file_bytes=500)


class TestExtractTextFromBytes:
    def test_plain_text(self) -> None:
        data = b"John Doe\nSoftware Engineer\n5 years experience"
        result = extract_text_from_bytes(data, "cv.txt")
        assert "John Doe" in result
        assert "Software Engineer" in result

    def test_unsupported_extension(self) -> None:
        with pytest.raises(CVProcessingError, match="Unsupported file type"):
            extract_text_from_bytes(b"data", "cv.xlsx")

    def test_empty_txt(self) -> None:
        result = extract_text_from_bytes(b"", "cv.txt")
        assert result == ""

    def test_docx_extracts_paragraphs(self) -> None:
        doc = Document()
        doc.add_paragraph("John Doe")
        doc.add_paragraph("Senior Engineer")
        buffer = io.BytesIO()
        doc.save(buffer)
        result = extract_text_from_bytes(buffer.getvalue(), "cv.docx")
        assert "John Doe" in result


class TestCleanCvText:
    def test_normalizes_whitespace(self) -> None:
        text = "John   Doe\n\n\n\n\nEngineer"
        result = clean_cv_text(text)
        assert "John Doe" in result
        assert "\n\n\n" not in result

    def test_removes_non_printable(self) -> None:
        text = "John\x00Doe\x01Engineer"
        result = clean_cv_text(text)
        assert "\x00" not in result
        assert "\x01" not in result

    def test_preserves_normal_content(self) -> None:
        text = "Senior Python Engineer\n- 5 years experience\n- Machine Learning"
        result = clean_cv_text(text)
        assert "Senior Python Engineer" in result
        assert "Machine Learning" in result


class TestTruncateToTokenLimit:
    def test_short_text_unchanged(self) -> None:
        text = "Short CV text"
        result = truncate_to_token_limit(text, max_tokens=100)
        assert result == text

    def test_long_text_truncated(self) -> None:
        text = "word " * 5000
        result = truncate_to_token_limit(text, max_tokens=50)
        assert "[CV truncated" in result
        assert len(result) < len(text)


class TestProcessCv:
    def test_full_pipeline_txt(self) -> None:
        cv = b"John Doe\n\n\nSenior Engineer\n\n\n\nPython, ML, FastAPI"
        result = process_cv(cv, "cv.txt")
        assert "John Doe" in result
        assert "Python" in result
        assert "\n\n\n" not in result

    def test_empty_cv_raises(self) -> None:
        with pytest.raises(CVProcessingError, match="empty"):
            process_cv(b"   \n   ", "cv.txt")

    def test_unsupported_type_raises(self) -> None:
        with pytest.raises(CVProcessingError, match="Unsupported"):
            process_cv(b"data", "cv.jpg")

    def test_oversized_cv_raises(self) -> None:
        with pytest.raises(CVProcessingError, match="too large"):
            process_cv(b"x" * 1000, "cv.txt", max_file_bytes=500)
